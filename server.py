"""
CaseCommand Server — FastAPI app on port 3000

Unified gateway for web dashboard, Telegram, WhatsApp, and SMS channels.
All channels route through the same Casey agent (src/agent.py).
"""

import os
import re
import json
import uuid
import logging
import io
from pathlib import Path
from datetime import datetime, timezone
from dotenv import load_dotenv

load_dotenv()

from fastapi import FastAPI, HTTPException, Depends, Security, Request, BackgroundTasks, UploadFile, File
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from supabase import create_client, Client
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uvicorn

from src.api_client import CaseCommandAI
from src.agent import process_message
from src.channels import telegram as tg_channel
from src.channels import twilio as tw_channel
from src.self_healer import start_healer_background

logger = logging.getLogger(__name__)

app = FastAPI(title="CaseCommand API")

SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_SECRET_KEY = os.environ.get("SUPABASE_SECRET_KEY", "") or os.environ.get("SUPABASE_SERVICE_KEY", "") or os.environ.get("SUPABASE_KEY", "")
AUTH_TOKEN = os.environ.get("AUTH_TOKEN", "")

supabase: Client = create_client(SUPABASE_URL, SUPABASE_SECRET_KEY)
ai_client = CaseCommandAI()

# ---------------------------------------------------------------------------
# Document storage
# ---------------------------------------------------------------------------

DOCUMENTS_DIR = Path(__file__).parent / "documents"
DOCUMENTS_DIR.mkdir(exist_ok=True)

OUTLINES_DIR = Path(__file__).parent / "outlines"
OUTLINES_DIR.mkdir(exist_ok=True)

# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------

security = HTTPBearer(auto_error=False)


def verify_token(credentials: HTTPAuthorizationCredentials = Security(security)):
    if not AUTH_TOKEN:
        return
    if not credentials or credentials.credentials != AUTH_TOKEN:
        raise HTTPException(status_code=401, detail="Invalid or missing token")


# ---------------------------------------------------------------------------
# Pydantic models
# ---------------------------------------------------------------------------


class CaseCreate(BaseModel):
    name: str
    type: str = "Other"
    client: str = "TBD"
    opposing: str = "TBD"


class ChatRequest(BaseModel):
    message: str
    session_id: str | None = None
    current_case_id: str | None = None


class AIRequest(BaseModel):
    system: str
    message: str
    max_tokens: int = 4096
    temperature: float = 0.3


class DiscoveryRequest(BaseModel):
    case_name: str
    discovery_type: str
    requests_and_responses: list


class SettlementRequest(BaseModel):
    case_name: str
    trigger_point: str
    valuation_data: dict
    recommendation_data: dict


class OutlineRequest(BaseModel):
    case_name: str
    witness_name: str
    witness_role: str = "witness"
    exam_type: str = "cross"
    case_theory: str = ""
    documents: list = []


# ---------------------------------------------------------------------------
# Document generation helpers
# ---------------------------------------------------------------------------

DOCUMENT_TRIGGER_WORDS = [
    "draft", "write", "prepare", "create a", "generate a", "compose",
    "produce", "draw up",
]
DOCUMENT_TYPE_WORDS = [
    "letter", "motion", "complaint", "demand", "agreement", "memo",
    "memorandum", "notice", "brief", "contract", "pleading", "declaration",
    "affidavit", "stipulation", "subpoena", "order", "outline", "response",
    "meet and confer", "m&c",
]

DOCUMENT_FORMAT_INSTRUCTIONS = (
    "\n\nIMPORTANT: When drafting any legal document (letter, motion, memo, "
    "demand letter, complaint, agreement, etc.), begin your response with this "
    "exact line:\n"
    "DOCUMENT_TITLE: [short descriptive title, e.g. Meet_and_Confer_Letter_Rodriguez]\n"
    "Then provide the full document text. This allows the system to automatically "
    "create a downloadable Word file for the attorney.\n\n"
)


def is_document_request(message: str) -> bool:
    msg = message.lower()
    has_action = any(kw in msg for kw in DOCUMENT_TRIGGER_WORDS)
    has_doc_type = any(dt in msg for dt in DOCUMENT_TYPE_WORDS)
    return has_action and has_doc_type


def _title_from_message(message: str) -> str:
    """Derive a safe document filename from the user's request."""
    skip = {"a", "an", "the", "for", "of", "in", "to", "me", "us", "please", "can", "you"}
    words = [w.capitalize() for w in message.split() if w.lower() not in skip]
    return "_".join(words[:6]) or "Legal_Document"


def extract_title_and_body(text: str, fallback: str = "Legal_Document") -> tuple[str, str]:
    """Return (title, body) — strips the DOCUMENT_TITLE line wherever it appears."""
    lines = text.strip().split("\n")
    for i, line in enumerate(lines):
        if line.strip().startswith("DOCUMENT_TITLE:"):
            title = line.replace("DOCUMENT_TITLE:", "").strip()
            body = "\n".join(lines[:i] + lines[i + 1:]).lstrip("\n")
            return title, body
    return fallback, text


def _add_run_with_text(para, text: str):
    """Add a run to a paragraph, applying bold for **...**-wrapped segments."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def build_docx(title: str, body: str) -> str:
    """Convert title + body text into a formatted .docx file. Returns filename."""
    doc = DocxDocument()

    # Page margins: 1.25" left/right, 1" top/bottom (standard legal)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.replace("_", " "))
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)
    doc.add_paragraph()  # blank spacer

    # Body — handle markdown-style headings and bold
    for line in body.split("\n"):
        stripped = line.rstrip()

        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            for run in h.runs:
                run.font.name = "Times New Roman"
            continue

        if stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            for run in h.runs:
                run.font.name = "Times New Roman"
            continue

        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            for run in h.runs:
                run.font.name = "Times New Roman"
            continue

        para = doc.add_paragraph()
        _add_run_with_text(para, stripped)

    safe = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "_")[:50]
    filename = f"{safe}_{uuid.uuid4().hex[:6]}.docx"
    doc.save(str(DOCUMENTS_DIR / filename))
    return filename


# ---------------------------------------------------------------------------
# Frontend
# ---------------------------------------------------------------------------


@app.get("/", include_in_schema=False)
def serve_frontend():
    index_path = Path(__file__).parent / "index.html"
    if index_path.exists():
        return HTMLResponse(content=index_path.read_text(encoding="utf-8"))
    return HTMLResponse(
        "<h1>CaseCommand</h1><p>Visit <a href='/docs'>/docs</a> for the API.</p>"
    )


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------


@app.get("/api/health")
def health(_: None = Depends(verify_token)):
    return {"status": "ok", "model": "claude-sonnet-4-6"}


@app.post("/api/ai")
async def ai_generate(req: AIRequest, _: None = Depends(verify_token)):
    """General-purpose AI endpoint for module features (Meet & Confer, Discovery, etc.)."""
    try:
        response = ai_client._call_api(req.system, req.message, max_tokens=req.max_tokens)
        return {"text": response["text"], "success": True, "usage": response.get("usage", {})}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/cases")
def list_cases(_: None = Depends(verify_token)):
    try:
        result = supabase.table("cases").select("*").order("created_at", desc=True).execute()
        return result.data or []
    except Exception as e:
        logger.error(f"list_cases DB error: {e}")
        raise HTTPException(status_code=500, detail=f"Database error listing cases: {str(e)}")


@app.get("/api/cases/{case_id}")
def get_case(case_id: str, _: None = Depends(verify_token)):
    try:
        result = supabase.table("cases").select("*").eq("id", case_id).execute()
        if not result.data:
            raise HTTPException(status_code=404, detail="Case not found")
        return result.data[0]
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"get_case DB error: {e}")
        raise HTTPException(status_code=500, detail=f"Database error fetching case: {str(e)}")


@app.post("/api/cases", status_code=201)
def create_case(case: CaseCreate, _: None = Depends(verify_token)):
    try:
        data = {
            "case_name": case.name,
            "case_type": case.type or "Other",
            "client_name": case.client or "TBD",
            "opposing_party": case.opposing or "TBD",
            "status": "active",
        }
        result = supabase.table("cases").insert(data).execute()
        if not result.data:
            logger.error("Case insert returned empty data — likely an RLS policy violation. "
                         "Ensure SUPABASE_SECRET_KEY is set to the service_role (sb_secret_*) key, "
                         "not the anon/publishable key.")
            raise HTTPException(
                status_code=500,
                detail="Case could not be saved. Database insert returned no data. "
                       "Check that SUPABASE_SECRET_KEY is the service_role key and that "
                       "the migration 002_fix_rls.sql has been run in Supabase."
            )
        return result.data[0]
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"create_case DB error: {e}")
        raise HTTPException(status_code=500, detail=f"Database error creating case: {str(e)}")


@app.get("/api/debug")
def debug_db(_: None = Depends(verify_token)):
    """Diagnostic endpoint — checks Supabase connectivity and RLS access."""
    report = {"supabase_url": SUPABASE_URL[:40] + "..." if SUPABASE_URL else "NOT SET",
              "key_set": bool(SUPABASE_SECRET_KEY),
              "key_prefix": SUPABASE_SECRET_KEY[:12] + "..." if SUPABASE_SECRET_KEY else "NOT SET"}
    try:
        result = supabase.table("cases").select("id,case_name,status").limit(5).execute()
        report["cases_readable"] = True
        report["cases_count"] = len(result.data or [])
        report["sample_cases"] = [c.get("case_name") for c in (result.data or [])]
    except Exception as e:
        report["cases_readable"] = False
        report["cases_error"] = str(e)
    try:
        test_data = {"case_name": "__debug_test__", "case_type": "Other",
                     "client_name": "TBD", "opposing_party": "TBD", "status": "active"}
        ins = supabase.table("cases").insert(test_data).execute()
        if ins.data:
            supabase.table("cases").delete().eq("case_name", "__debug_test__").execute()
            report["cases_writable"] = True
        else:
            report["cases_writable"] = False
            report["write_error"] = "Insert returned empty data — RLS policy blocking write"
    except Exception as e:
        report["cases_writable"] = False
        report["write_error"] = str(e)
    return report


@app.post("/api/chat")
async def chat(req: ChatRequest, _: None = Depends(verify_token)):
    """Web dashboard chat — routes through the unified Casey agent with persistent memory."""
    session_id = req.session_id or str(uuid.uuid4())
    ctx = {"supabase": supabase, "session_id": session_id}

    # Load conversation history from database for this session
    history = _load_conversation_history(session_id)

    try:
        response = await process_message(req.message, conversation_history=history, context=ctx)
    except Exception as e:
        logger.error(f"process_message failed: {e}")
        raise HTTPException(status_code=500, detail=f"Agent error: {str(e)}")

    reply_text, document = _handle_agent_response(response, req.message)

    # Persist this exchange to the database
    _save_conversation_messages(session_id, "web", req.message, reply_text, response)

    return {
        "response": reply_text,
        "model": response["model"],
        "document": document,
        "session_id": session_id,
    }


def _load_conversation_history(session_id: str, limit: int = 20) -> list:
    """Load recent conversation messages from DB for a session."""
    try:
        result = (
            supabase.table("conversation_messages")
            .select("role,content")
            .eq("session_id", session_id)
            .order("created_at", desc=False)
            .limit(limit)
            .execute()
        )
        return [{"role": r["role"], "content": r["content"]} for r in (result.data or [])]
    except Exception as e:
        logger.warning(f"Could not load conversation history: {e}")
        return []


def _save_conversation_messages(
    session_id: str, channel: str, user_msg: str, assistant_reply: str, response: dict
):
    """Persist a user+assistant message pair to the database."""
    try:
        now = datetime.now(timezone.utc).isoformat()
        rows = [
            {
                "session_id": session_id,
                "channel": channel,
                "role": "user",
                "content": user_msg,
                "metadata": {},
                "created_at": now,
            },
            {
                "session_id": session_id,
                "channel": channel,
                "role": "assistant",
                "content": assistant_reply,
                "metadata": {
                    "model": response.get("model", ""),
                    "tool_calls": [tc.get("name") for tc in response.get("tool_calls", [])],
                },
                "created_at": now,
            },
        ]
        supabase.table("conversation_messages").insert(rows).execute()
    except Exception as e:
        logger.warning(f"Could not save conversation messages: {e}")


# ---------------------------------------------------------------------------
# Document analysis for new case creation
# ---------------------------------------------------------------------------


def _extract_text_from_upload(content: bytes, filename: str) -> str:
    """Extract plain text from an uploaded file (txt, docx, pdf)."""
    name_lower = filename.lower()

    if name_lower.endswith(".txt"):
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            return content.decode("latin-1", errors="replace")

    if name_lower.endswith(".docx") or name_lower.endswith(".doc"):
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.warning(f"Could not parse docx: {e}")
            return ""

    if name_lower.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pages = []
            for page in reader.pages[:10]:  # First 10 pages
                pages.append(page.extract_text() or "")
            return "\n".join(pages)
        except Exception as e:
            logger.warning(f"Could not parse pdf: {e}")
            return ""

    # Fallback: try as UTF-8 text
    try:
        return content.decode("utf-8", errors="replace")
    except Exception:
        return ""


async def _call_claude_for_case_extraction(text: str) -> dict:
    """Ask Claude to extract case metadata from document text."""
    import httpx as _httpx
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    model = os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6")

    system = (
        "You are a legal document analyst. Given document text, extract case information. "
        "Respond ONLY with a JSON object (no markdown, no explanation) with these fields:\n"
        '{"case_name": "Plaintiff v. Defendant", "case_type": "PI|Employment|Contract|Criminal|Family|Real Estate|Other", '
        '"client_name": "full name or TBD", "opposing_party": "full name or TBD", '
        '"summary": "1-2 sentence case summary or empty string"}\n'
        "If you cannot determine a value, use TBD. For case_name, construct it as 'Client v. Opposing Party' if possible."
    )
    prompt = f"Extract case information from this document:\n\n{text[:4000]}"

    payload = {
        "model": model,
        "max_tokens": 512,
        "system": system,
        "messages": [{"role": "user", "content": prompt}],
    }
    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    _fallback = {
        "case_name": "New Case",
        "case_type": "Other",
        "client_name": "TBD",
        "opposing_party": "TBD",
        "summary": "",
    }
    try:
        async with _httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post("https://api.anthropic.com/v1/messages", headers=headers, json=payload)
            resp.raise_for_status()
            data = resp.json()
    except Exception as e:
        logger.warning(f"Claude API call failed during case extraction: {e}")
        return _fallback

    raw = ""
    for block in data.get("content", []):
        if block.get("type") == "text":
            raw += block["text"]

    # Strip markdown code fences if present
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return _fallback


@app.post("/api/analyze-document")
async def analyze_document(file: UploadFile = File(...), _: None = Depends(verify_token)):
    """
    Upload a case document (PDF, DOCX, TXT). Returns extracted case metadata
    for human-in-the-loop review before case creation.
    """
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Empty file")

    text = _extract_text_from_upload(content, file.filename or "")
    if not text.strip():
        # Return a graceful fallback so the UI still shows the confirmation form
        return {
            "case_name": "New Case",
            "case_type": "Other",
            "client_name": "TBD",
            "opposing_party": "TBD",
            "summary": "Could not extract text from document. Please fill in the details manually.",
            "filename": file.filename,
        }

    extracted = await _call_claude_for_case_extraction(text)
    extracted["filename"] = file.filename
    return extracted


@app.get("/api/documents/{filename}", include_in_schema=False)
def download_document(filename: str):
    # Prevent path traversal
    safe_filename = Path(filename).name
    filepath = DOCUMENTS_DIR / safe_filename
    if not filepath.exists() or not filepath.is_file():
        raise HTTPException(status_code=404, detail="Document not found")
    return FileResponse(
        str(filepath),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=safe_filename,
    )


@app.get("/api/deadlines")
def deadlines(_: None = Depends(verify_token)):
    result = (
        supabase.table("cases")
        .select("*")
        .eq("status", "active")
        .order("created_at", desc=False)
        .execute()
    )
    return result.data


@app.post("/api/discovery")
def analyze_discovery(req: DiscoveryRequest, _: None = Depends(verify_token)):
    response = ai_client.analyze_discovery_responses(
        case_name=req.case_name,
        discovery_type=req.discovery_type,
        requests_and_responses=req.requests_and_responses,
    )
    return {"result": response["text"], "model": response["model"]}


@app.post("/api/settlement")
def generate_settlement(req: SettlementRequest, _: None = Depends(verify_token)):
    response = ai_client.generate_settlement_narrative(
        case_name=req.case_name,
        trigger_point=req.trigger_point,
        valuation_data=req.valuation_data,
        recommendation_data=req.recommendation_data,
    )
    return {"result": response["text"], "model": response["model"]}


# ---------------------------------------------------------------------------
# TrialOutline Pro — landscape HTML viewer
# ---------------------------------------------------------------------------


def build_outline_html(case_name: str, witness_name: str, exam_type: str, outline_text: str) -> str:
    """Parse AI outline text into a self-contained landscape HTML viewer."""
    exam_label = "CROSS-EXAMINATION" if exam_type.lower() == "cross" else "DIRECT EXAMINATION"

    # Parse sections and questions
    sections: list[dict] = []
    current: dict | None = None
    for raw in outline_text.split("\n"):
        line = raw.strip()
        if not line:
            continue
        if re.match(r"^(#{1,3}|[IVXLC]+\.)\s", line) or (re.match(r"^\d+\.", line) and len(line) < 70 and line[0].isdigit() and not line[0:2].isdigit()):
            title = re.sub(r"^#{1,3}\s*|^[IVXLC]+\.\s*", "", line).strip()
            current = {"title": title, "goal": "", "questions": []}
            sections.append(current)
        elif line.lower().startswith("goal:") and current is not None:
            current["goal"] = line[5:].strip()
        elif current is not None and re.match(r"^\d+\.", line):
            q = re.sub(r"^\d+\.\s*", "", line)
            if q:
                current["questions"].append(q)
        elif current is not None and re.match(r"^[-•Q]\s", line):
            q = re.sub(r"^[-•Q][:.]?\s*", "", line)
            if q:
                current["questions"].append(q)

    if not sections:
        questions = [re.sub(r"^\d+\.\s*|^[-•]\s*", "", l.strip())
                     for l in outline_text.split("\n")
                     if l.strip() and re.match(r"^\d+\.|^[-•]", l.strip())]
        sections = [{"title": "Examination Questions", "goal": "", "questions": questions or ["(No questions parsed)"]}]

    q_num = 1
    sections_html = ""
    for sec in sections:
        items_html = ""
        for q in sec["questions"]:
            items_html += f'<div class="q-item" data-q="{q_num}" onclick="selectQ(this)">{q_num}. {q}</div>\n'
            q_num += 1
        goal_html = f'<div class="sec-goal">Goal: {sec["goal"]}</div>' if sec["goal"] else ""
        sections_html += f"""<div class="section">
  <div class="sec-title">{sec["title"].upper()}</div>{goal_html}{items_html}</div>"""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{exam_label} — {witness_name}</title>
<style>
@page {{ size: landscape; margin: 0.5in; }}
*,*::before,*::after {{ box-sizing:border-box; margin:0; padding:0; }}
body {{ font-family:"Times New Roman",serif; background:#f0f2f5; height:100vh; display:flex; flex-direction:column; overflow:hidden; }}
.hdr {{ background:#1a1a2e; color:#fff; padding:10px 20px; display:flex; align-items:center; justify-content:space-between; flex-shrink:0; }}
.hdr-title {{ font-size:15px; font-weight:700; }}
.hdr-sub {{ font-size:11px; color:#aaa; margin-top:2px; }}
.hdr-btns {{ display:flex; gap:8px; }}
.hdr-btns button {{ padding:4px 12px; border-radius:4px; border:1px solid #555; background:transparent; color:#ccc; cursor:pointer; font-size:12px; }}
.hdr-btns button:hover {{ background:#2a2a4e; color:#fff; }}
.workspace {{ display:flex; flex:1; overflow:hidden; }}
.q-panel {{ width:42%; background:#fff; border-right:2px solid #dee2e6; display:flex; flex-direction:column; overflow:hidden; }}
.q-panel-hdr {{ background:#2c3e50; color:#fff; padding:8px 14px; font-size:12px; font-weight:700; letter-spacing:.3px; flex-shrink:0; }}
.q-scroll {{ flex:1; overflow-y:auto; padding:10px; }}
.section {{ margin-bottom:14px; }}
.sec-title {{ font-size:10px; font-weight:700; color:#495057; background:#e9ecef; padding:3px 8px; border-radius:3px; margin-bottom:3px; letter-spacing:.5px; }}
.sec-goal {{ font-size:10px; color:#6c757d; font-style:italic; padding:2px 8px 3px; }}
.q-item {{ padding:6px 10px; border-radius:4px; margin-bottom:2px; font-size:12px; cursor:pointer; border-left:3px solid transparent; line-height:1.4; transition:background .1s; }}
.q-item:hover {{ background:#f0f4ff; border-left-color:#4a90e2; }}
.q-item.active {{ background:#dbeafe; border-left-color:#1d4ed8; font-weight:600; }}
.kb-hint {{ font-size:10px; color:#6c757d; text-align:center; padding:5px; border-top:1px solid #dee2e6; background:#f8f9fa; flex-shrink:0; }}
.ex-panel {{ width:58%; background:#fff; display:flex; flex-direction:column; overflow:hidden; }}
.ex-panel-hdr {{ background:#155724; color:#fff; padding:8px 14px; font-size:12px; font-weight:700; flex-shrink:0; }}
.ex-content {{ flex:1; overflow-y:auto; padding:20px; }}
.q-display {{ background:#1d4ed8; color:#fff; padding:14px 18px; border-radius:8px; margin-bottom:20px; font-size:14px; font-weight:500; line-height:1.5; display:none; }}
.q-display.show {{ display:block; }}
.exhibit-ph {{ border:2px dashed #dee2e6; border-radius:8px; padding:40px 30px; text-align:center; color:#adb5bd; }}
.exhibit-ph h3 {{ font-size:15px; margin-bottom:8px; color:#6c757d; }}
.exhibit-ph p {{ font-size:12px; line-height:1.6; }}
@media print {{
  body {{ height:auto; overflow:visible; }}
  .hdr-btns {{ display:none; }}
  .workspace {{ page-break-inside:avoid; }}
}}
</style>
</head>
<body>
<div class="hdr">
  <div>
    <div class="hdr-title">⚖️ {exam_label}: {witness_name}</div>
    <div class="hdr-sub">{case_name} &nbsp;·&nbsp; CaseCommand by LawClaw</div>
  </div>
  <div class="hdr-btns">
    <button onclick="prev()">← Prev</button>
    <span id="counter" style="color:#aaa;font-size:12px;align-self:center">0 / {q_num - 1}</span>
    <button onclick="next()">Next →</button>
    <button onclick="window.print()">🖨 Print</button>
    <button onclick="window.close()">✕ Close</button>
  </div>
</div>
<div class="workspace">
  <div class="q-panel">
    <div class="q-panel-hdr">📋 Examination Questions</div>
    <div class="q-scroll" id="qScroll">{sections_html}</div>
    <div class="kb-hint">↑ ↓ Arrow keys to navigate · Click question to select</div>
  </div>
  <div class="ex-panel">
    <div class="ex-panel-hdr">📄 Current Question / Exhibit Reference</div>
    <div class="ex-content">
      <div class="q-display" id="qDisplay"></div>
      <div class="exhibit-ph" id="exPh">
        <h3>Select a question to begin</h3>
        <p>Click any question on the left or use ↑↓ arrow keys.<br>
        The question will appear here in large text for easy reference at counsel table.<br><br>
        Exhibits can be added to CaseCommand cases to display alongside questions.</p>
      </div>
    </div>
  </div>
</div>
<script>
const items = Array.from(document.querySelectorAll('.q-item'));
const total = items.length;
let idx = -1;
function selectQ(el) {{
  items.forEach(i => i.classList.remove('active'));
  el.classList.add('active');
  idx = items.indexOf(el);
  document.getElementById('qDisplay').textContent = el.textContent;
  document.getElementById('qDisplay').classList.add('show');
  document.getElementById('exPh').style.display = 'none';
  document.getElementById('counter').textContent = (idx+1) + ' / ' + total;
  el.scrollIntoView({{block:'nearest'}});
}}
function next() {{ if (idx < total-1) selectQ(items[idx+1]); }}
function prev() {{ if (idx > 0) selectQ(items[idx-1]); }}
document.addEventListener('keydown', e => {{
  if (e.key==='ArrowDown') {{ e.preventDefault(); next(); }}
  if (e.key==='ArrowUp')   {{ e.preventDefault(); prev(); }}
}});
if (total > 0) selectQ(items[0]);
</script>
</body>
</html>"""


@app.post("/api/outline")
def generate_outline(req: OutlineRequest, _: None = Depends(verify_token)):
    response = ai_client.generate_examination_outline(
        case_name=req.case_name,
        witness_name=req.witness_name,
        witness_role=req.witness_role,
        exam_type=req.exam_type,
        case_documents=req.documents,
        case_theory=req.case_theory,
    )
    html = build_outline_html(
        case_name=req.case_name,
        witness_name=req.witness_name,
        exam_type=req.exam_type,
        outline_text=response["text"],
    )
    outline_id = uuid.uuid4().hex[:8]
    filename = f"outline_{outline_id}.html"
    (OUTLINES_DIR / filename).write_text(html, encoding="utf-8")
    return {"url": f"/api/outlines/{filename}", "model": response["model"]}


@app.get("/api/outlines/{filename}", include_in_schema=False)
def serve_outline(filename: str):
    safe = Path(filename).name
    filepath = OUTLINES_DIR / safe
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="Outline not found")
    return HTMLResponse(content=filepath.read_text(encoding="utf-8"))


# ---------------------------------------------------------------------------
# Shared agent response handler (used by web chat + messaging channels)
# ---------------------------------------------------------------------------


def _handle_agent_response(response: dict, original_message: str) -> tuple[str, dict | None]:
    """
    Post-process agent response: build .docx if a document tool was called,
    handle DOCUMENT_TITLE: fallback, append download link.

    Returns (reply_text, document_dict_or_None).
    """
    reply_text = response["reply"].strip()
    document = None

    for tool_call in response.get("tool_calls", []):
        if tool_call["name"] == "generate_legal_document":
            try:
                title = tool_call["input"]["title"]
                body = tool_call["input"]["body"]
                filename = build_docx(title, body)
                document = {
                    "filename": filename,
                    "url": f"/api/documents/{filename}",
                    "title": title.replace("_", " "),
                }
                if not reply_text:
                    reply_text = body
            except Exception:
                pass

    if not document and "DOCUMENT_TITLE:" in reply_text:
        try:
            fallback = _title_from_message(original_message)
            title, body = extract_title_and_body(reply_text, fallback=fallback)
            filename = build_docx(title, body)
            document = {
                "filename": filename,
                "url": f"/api/documents/{filename}",
                "title": title.replace("_", " "),
            }
            reply_text = body
        except Exception:
            pass

    if document:
        reply_text += (
            f"\n\n📄 **Document Ready:** {document['title']}\n"
            f"Download: {document['url']}"
        )

    return reply_text, document


# ---------------------------------------------------------------------------
# Telegram webhook
# ---------------------------------------------------------------------------


@app.post("/webhook/telegram")
async def telegram_webhook(request: Request, background_tasks: BackgroundTasks):
    """
    Receive Telegram updates via webhook. Processes messages through
    the Casey agent and replies in-channel.
    """
    if not tg_channel.is_configured():
        raise HTTPException(status_code=503, detail="Telegram not configured")

    try:
        update = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON")

    parsed = tg_channel.parse_update(update)
    if not parsed:
        return {"ok": True}  # Non-message update (e.g., button callback)

    if not tg_channel.is_authorized(parsed["user_id"]):
        await tg_channel.send_message(
            parsed["chat_id"],
            "⛔ You are not authorized to use CaseCommand. "
            "Contact your administrator to get access.",
            reply_to=parsed["message_id"],
        )
        return {"ok": True}

    # Process in background so Telegram doesn't time out
    background_tasks.add_task(
        _process_telegram_message,
        parsed["chat_id"],
        parsed["text"],
        parsed["message_id"],
    )
    return {"ok": True}


async def _process_telegram_message(chat_id: int, text: str, reply_to: int):
    """Background task: run agent and reply on Telegram."""
    try:
        session_id = f"telegram_{chat_id}"
        ctx = {"supabase": supabase, "session_id": session_id}
        history = _load_conversation_history(session_id, limit=20)
        response = await process_message(text, conversation_history=history, context=ctx)
        reply_text, document = _handle_agent_response(response, text)
        _save_conversation_messages(session_id, "telegram", text, reply_text, response)

        # Send text reply
        await tg_channel.send_message(chat_id, reply_text, reply_to=reply_to)

        # If a document was generated, also send the file
        if document:
            filepath = DOCUMENTS_DIR / document["filename"]
            if filepath.exists():
                await tg_channel.send_document(
                    chat_id,
                    str(filepath),
                    caption=f"📄 {document['title']}",
                )
    except Exception as e:
        logger.error(f"Telegram processing error: {e}")
        await tg_channel.send_message(
            chat_id,
            "⚠️ Sorry, I encountered an error processing your request. Please try again.",
            reply_to=reply_to,
        )


# ---------------------------------------------------------------------------
# Twilio webhook (SMS + WhatsApp)
# ---------------------------------------------------------------------------


@app.post("/webhook/twilio")
async def twilio_webhook(request: Request, background_tasks: BackgroundTasks):
    """
    Receive Twilio SMS/WhatsApp webhook callbacks. Processes messages
    through the Casey agent and replies via Twilio API.
    """
    if not tw_channel.is_configured():
        raise HTTPException(status_code=503, detail="Twilio not configured")

    try:
        form = await request.form()
        form_data = dict(form)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid form data")

    parsed = tw_channel.parse_webhook(form_data)
    if not parsed:
        return {"ok": True}

    if not tw_channel.is_authorized(parsed["from"]):
        await tw_channel.send_message(
            parsed["from"],
            "⛔ You are not authorized to use CaseCommand.",
            channel=parsed["channel"],
        )
        # Return TwiML empty response
        return HTMLResponse(
            content="<Response></Response>",
            media_type="application/xml",
        )

    # Process in background
    background_tasks.add_task(
        _process_twilio_message,
        parsed["from"],
        parsed["body"],
        parsed["channel"],
    )

    # Return empty TwiML so Twilio doesn't retry
    return HTMLResponse(
        content="<Response></Response>",
        media_type="application/xml",
    )


async def _process_twilio_message(from_number: str, text: str, channel: str):
    """Background task: run agent and reply via Twilio SMS/WhatsApp."""
    try:
        session_id = f"twilio_{from_number.replace('+', '').replace(':', '_')}"
        ctx = {"supabase": supabase, "session_id": session_id}
        history = _load_conversation_history(session_id, limit=20)
        response = await process_message(text, conversation_history=history, context=ctx)
        reply_text, document = _handle_agent_response(response, text)
        _save_conversation_messages(session_id, channel, text, reply_text, response)

        # For messaging channels, include download URL as full URL if document generated
        if document:
            base_url = os.environ.get("BASE_URL", "").rstrip("/")
            if base_url:
                reply_text = reply_text.replace(
                    document["url"],
                    f"{base_url}{document['url']}",
                )

        await tw_channel.send_message(from_number, reply_text, channel=channel)
    except Exception as e:
        logger.error(f"Twilio processing error: {e}")
        await tw_channel.send_message(
            from_number,
            "⚠️ Sorry, I encountered an error. Please try again.",
            channel=channel,
        )


# ---------------------------------------------------------------------------
# Channel setup helpers
# ---------------------------------------------------------------------------


@app.post("/api/channels/telegram/setup")
async def setup_telegram(_: None = Depends(verify_token)):
    """Register the Telegram webhook. Call once after deployment."""
    if not tg_channel.is_configured():
        raise HTTPException(status_code=400, detail="TELEGRAM_BOT_TOKEN not set")
    base_url = os.environ.get("BASE_URL", "").rstrip("/")
    if not base_url:
        raise HTTPException(
            status_code=400,
            detail="BASE_URL env var not set. Set it to your public URL (e.g., https://casecommand.onrender.com)",
        )
    result = await tg_channel.set_webhook(f"{base_url}/webhook/telegram")
    return result


@app.get("/api/channels/status")
async def channel_status(_: None = Depends(verify_token)):
    """Check which messaging channels are configured, with live webhook status."""
    tg_webhook = None
    tg_bot = None
    if tg_channel.is_configured():
        tg_webhook = await tg_channel.get_webhook_info()
        tg_bot = await tg_channel.get_bot_info()

    webhook_url = ""
    webhook_active = False
    if tg_webhook and tg_webhook.get("ok"):
        result = tg_webhook.get("result", {})
        webhook_url = result.get("url", "")
        webhook_active = bool(webhook_url)

    bot_username = ""
    if tg_bot and tg_bot.get("ok"):
        bot_username = tg_bot.get("result", {}).get("username", "")

    return {
        "telegram": {
            "configured": tg_channel.is_configured(),
            "webhook": "/webhook/telegram",
            "webhook_url": webhook_url,
            "webhook_active": webhook_active,
            "bot_username": bot_username,
        },
        "twilio_sms": {
            "configured": tw_channel.is_configured() and bool(tw_channel.TWILIO_PHONE_NUMBER),
            "webhook": "/webhook/twilio",
        },
        "twilio_whatsapp": {
            "configured": tw_channel.is_configured() and bool(tw_channel.TWILIO_WHATSAPP_NUMBER),
            "webhook": "/webhook/twilio",
        },
        "web": {
            "configured": True,
            "endpoint": "/api/chat",
        },
    }


# ---------------------------------------------------------------------------
# Agent Lab — Admin endpoints for reviewing agent outputs
# ---------------------------------------------------------------------------


class AgentOutputUpdate(BaseModel):
    status: str  # "applied" or "dismissed"


@app.get("/api/agent-outputs")
def list_agent_outputs(
    status: str | None = None,
    agent: str | None = None,
    limit: int = 50,
    _: None = Depends(verify_token),
):
    """List agent outputs, optionally filtered by status or agent name."""
    query = supabase.table("agent_outputs").select("*").order("created_at", desc=True).limit(limit)
    if status:
        query = query.eq("status", status)
    if agent:
        query = query.eq("agent_name", agent)
    result = query.execute()
    return result.data


@app.get("/api/agent-outputs/summary")
def agent_outputs_summary(_: None = Depends(verify_token)):
    """Get summary stats for the Agent Lab dashboard."""
    all_outputs = supabase.table("agent_outputs").select("id,agent_name,status,priority,run_id,created_at").order("created_at", desc=True).execute()
    data = all_outputs.data or []

    pending = [r for r in data if r["status"] == "pending"]
    applied = [r for r in data if r["status"] == "applied"]
    dismissed = [r for r in data if r["status"] == "dismissed"]

    # Group by agent
    agents = {}
    for r in data:
        name = r["agent_name"]
        if name not in agents:
            agents[name] = {"total": 0, "pending": 0, "applied": 0}
        agents[name]["total"] += 1
        if r["status"] == "pending":
            agents[name]["pending"] += 1
        elif r["status"] == "applied":
            agents[name]["applied"] += 1

    # Latest run
    run_ids = list({r["run_id"] for r in data if r.get("run_id")})
    latest_run = run_ids[0] if run_ids else None

    return {
        "total": len(data),
        "pending": len(pending),
        "applied": len(applied),
        "dismissed": len(dismissed),
        "agents": agents,
        "latest_run_id": latest_run,
    }


@app.patch("/api/agent-outputs/{output_id}")
def update_agent_output(output_id: str, update: AgentOutputUpdate, _: None = Depends(verify_token)):
    """Mark an agent output as applied or dismissed."""
    update_data = {"status": update.status}
    if update.status == "applied":
        from datetime import datetime, timezone
        update_data["applied_at"] = datetime.now(timezone.utc).isoformat()
        update_data["applied_by"] = "admin"

    result = supabase.table("agent_outputs").update(update_data).eq("id", output_id).execute()
    if not result.data:
        raise HTTPException(status_code=404, detail="Agent output not found")
    return result.data[0]


# ---------------------------------------------------------------------------
# Auto-setup / Self-healing — runs on startup and via /api/setup endpoint
# ---------------------------------------------------------------------------


def _run_auto_setup() -> dict:
    """
    Auto-healing startup routine.
    Tests DB connectivity and write access; logs actionable guidance if broken.
    Returns a status dict.
    """
    report = {"success": True, "issues": [], "fixes": []}

    # 1. Test connectivity
    try:
        result = supabase.table("cases").select("id").limit(1).execute()
        report["db_readable"] = True
        report["cases_count"] = len(result.data or [])
    except Exception as e:
        report["db_readable"] = False
        report["success"] = False
        report["issues"].append(f"Database not readable: {e}")
        logger.error(f"[AUTO-SETUP] DB read failed: {e}")
        return report

    # 2. Test write access
    try:
        test = {
            "case_name": "__startup_test__",
            "case_type": "Other",
            "client_name": "TBD",
            "opposing_party": "TBD",
            "status": "active",
        }
        ins = supabase.table("cases").insert(test).execute()
        if ins.data:
            # Clean up
            supabase.table("cases").delete().eq("case_name", "__startup_test__").execute()
            report["db_writable"] = True
            report["fixes"].append("Database write access confirmed — no RLS issues.")
        else:
            report["db_writable"] = False
            report["success"] = False
            report["issues"].append(
                "Database INSERT returned empty data. RLS policy is blocking writes. "
                "ACTION REQUIRED: Run database/migrations/002_fix_rls_backend_access.sql "
                "in the Supabase SQL Editor. Also confirm SUPABASE_SECRET_KEY is the "
                "service_role key (starts with 'eyJ' or 'sb_secret_')."
            )
            logger.error(
                "[AUTO-SETUP] DB write blocked by RLS. "
                "Run database/migrations/002_fix_rls_backend_access.sql in Supabase SQL Editor."
            )
    except Exception as e:
        report["db_writable"] = False
        report["success"] = False
        report["issues"].append(f"Database write test failed: {e}")
        logger.error(f"[AUTO-SETUP] DB write test error: {e}")

    # 3. Test casey_memory table
    try:
        supabase.table("casey_memory").select("key").limit(1).execute()
        report["memory_readable"] = True
    except Exception as e:
        report["memory_readable"] = False
        report["issues"].append(
            f"casey_memory table not found or not readable: {e}. "
            "Run database/migrations/001_casey_memory.sql in Supabase SQL Editor."
        )
        logger.warning(f"[AUTO-SETUP] casey_memory issue: {e}")

    if report["success"]:
        logger.info("[AUTO-SETUP] All systems operational. Database read/write confirmed.")
    else:
        logger.error(f"[AUTO-SETUP] Issues found: {report['issues']}")

    return report


@app.on_event("startup")
async def startup_event():
    """Run auto-setup diagnostics and start self-healer on server startup."""
    import asyncio
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, _run_auto_setup)
    # Start the continuous self-healing background task
    asyncio.create_task(_healer_task())


async def _healer_task():
    """Background self-healing loop."""
    from src.self_healer import healer_loop
    await healer_loop(supabase)


@app.get("/api/setup")
def run_setup(_: None = Depends(verify_token)):
    """
    Re-run the auto-setup diagnostic and return full status report.
    Call this after changing environment variables to verify DB connectivity.
    """
    return _run_auto_setup()


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=3000)

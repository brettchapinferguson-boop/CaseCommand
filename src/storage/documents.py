"""
CaseCommand — Document Storage

Handles document generation (.docx) and storage.
Supports both local filesystem (dev) and Supabase Storage (production).
"""

from __future__ import annotations

import re
import uuid
import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Local fallback directory (used in dev or when object storage unavailable)
LOCAL_DIR = Path(__file__).parent.parent.parent / "documents"
LOCAL_DIR.mkdir(exist_ok=True)

OUTLINES_DIR = Path(__file__).parent.parent.parent / "outlines"
OUTLINES_DIR.mkdir(exist_ok=True)

BUCKET_NAME = "documents"


class DocumentStore:
    """
    Manages document generation and storage.

    In production, uploads to Supabase Storage and returns signed URLs.
    Falls back to local filesystem in development.
    """

    def __init__(self, supabase_client=None):
        self.db = supabase_client
        self._bucket_verified = False

    async def _ensure_bucket(self):
        """Create the storage bucket if it doesn't exist."""
        if self._bucket_verified or not self.db:
            return
        try:
            self.db.storage.get_bucket(BUCKET_NAME)
            self._bucket_verified = True
        except Exception:
            try:
                self.db.storage.create_bucket(
                    BUCKET_NAME,
                    options={"public": False, "file_size_limit": 10485760},
                )
                self._bucket_verified = True
            except Exception as e:
                logger.warning("Could not create storage bucket: %s", e)

    def build_docx(self, title: str, body: str, org_id: str | None = None) -> tuple[str, str]:
        """
        Build a .docx file from title and body text.

        Returns (filename, local_path).
        """
        doc = DocxDocument()

        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title.replace("_", " "))
        title_run.bold = True
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(14)
        doc.add_paragraph()

        # Body
        for line in body.split("\n"):
            stripped = line.rstrip()

            if not stripped:
                doc.add_paragraph()
                continue

            if stripped.startswith("### "):
                h = doc.add_heading(stripped[4:], level=3)
                for run in h.runs:
                    run.font.name = "Times New Roman"
                continue

            if stripped.startswith("## "):
                h = doc.add_heading(stripped[3:], level=2)
                for run in h.runs:
                    run.font.name = "Times New Roman"
                continue

            if stripped.startswith("# "):
                h = doc.add_heading(stripped[2:], level=1)
                for run in h.runs:
                    run.font.name = "Times New Roman"
                continue

            para = doc.add_paragraph()
            self._add_run_with_text(para, stripped)

        safe = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "_")[:50]
        filename = f"{safe}_{uuid.uuid4().hex[:6]}.docx"

        # Use org-scoped path if available
        if org_id:
            subdir = LOCAL_DIR / org_id
            subdir.mkdir(exist_ok=True)
            local_path = str(subdir / filename)
        else:
            local_path = str(LOCAL_DIR / filename)

        doc.save(local_path)
        return filename, local_path

    async def upload_to_storage(self, filename: str, local_path: str, org_id: str) -> str | None:
        """
        Upload a document to Supabase Storage.

        Returns the storage path, or None if upload fails.
        """
        if not self.db:
            return None

        await self._ensure_bucket()

        storage_path = f"{org_id}/{filename}"
        try:
            with open(local_path, "rb") as f:
                self.db.storage.from_(BUCKET_NAME).upload(
                    storage_path,
                    f.read(),
                    file_options={
                        "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    },
                )
            return storage_path
        except Exception as e:
            logger.error("Storage upload failed: %s", e)
            return None

    def get_signed_url(self, storage_path: str, expires_in: int = 3600) -> str | None:
        """Generate a signed download URL (valid for expires_in seconds)."""
        if not self.db:
            return None
        try:
            result = self.db.storage.from_(BUCKET_NAME).create_signed_url(
                storage_path, expires_in
            )
            return result.get("signedURL")
        except Exception as e:
            logger.error("Signed URL generation failed: %s", e)
            return None

    @staticmethod
    def _add_run_with_text(para, text: str):
        """Add a run to a paragraph, applying bold for **...**-wrapped segments."""
        parts = re.split(r"(\*\*[^*]+\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def extract_title_and_body(text: str, fallback: str = "Legal_Document") -> tuple[str, str]:
    """Return (title, body) -- strips the DOCUMENT_TITLE line wherever it appears."""
    lines = text.strip().split("\n")
    for i, line in enumerate(lines):
        if line.strip().startswith("DOCUMENT_TITLE:"):
            title = line.replace("DOCUMENT_TITLE:", "").strip()
            body = "\n".join(lines[:i] + lines[i + 1:]).lstrip("\n")
            return title, body
    return fallback, text


def title_from_message(message: str) -> str:
    """Derive a safe document filename from the user's request."""
    skip = {"a", "an", "the", "for", "of", "in", "to", "me", "us", "please", "can", "you"}
    words = [w.capitalize() for w in message.split() if w.lower() not in skip]
    return "_".join(words[:6]) or "Legal_Document"
